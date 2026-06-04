# -*- coding: utf-8 -*-
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class Exporter:
    @staticmethod
    def export_to_excel(data, file_path):
        wb = Workbook()
        ws = wb.active
        ws.title = "大创项目数据"

        row = 1
        for section_title, section_data in data.items():
            ws.cell(row=row, column=1, value=section_title)
            ws.cell(row=row, column=1).font = ws.cell(row=row, column=1).font.copy(bold=True)
            row += 1

            if isinstance(section_data, list):
                for item in section_data:
                    if isinstance(item, dict):
                        for key, value in item.items():
                            ws.cell(row=row, column=1, value=key)
                            ws.cell(row=row, column=2, value=str(value))
                            row += 1
                    else:
                        ws.cell(row=row, column=1, value=str(item))
                        row += 1
            elif isinstance(section_data, dict):
                for key, value in section_data.items():
                    ws.cell(row=row, column=1, value=key)
                    ws.cell(row=row, column=2, value=str(value))
                    row += 1
            else:
                ws.cell(row=row, column=1, value=str(section_data))
                row += 1
            row += 1

        wb.save(file_path)

    @staticmethod
    def export_to_word(data, file_path):
        doc = Document()

        title = doc.add_heading("大创项目背景数据报告", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for section_title, section_data in data.items():
            doc.add_heading(section_title, level=1)

            if isinstance(section_data, list):
                for item in section_data:
                    if isinstance(item, dict):
                        for key, value in item.items():
                            p = doc.add_paragraph()
                            p.add_run(f"{key}: ").bold = True
                            p.add_run(str(value))
                    else:
                        doc.add_paragraph(str(item))
            elif isinstance(section_data, dict):
                for key, value in section_data.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{key}: ").bold = True
                    p.add_run(str(value))
            else:
                doc.add_paragraph(str(section_data))

        doc.save(file_path)

